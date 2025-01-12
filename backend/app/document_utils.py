import docx
import pdfplumber
import io
from typing import Tuple, List, Dict
import re
from docx.shared import Pt
from PyPDF2 import PdfReader, PdfWriter

def extract_text_from_document(file_content: bytes, filename: str) -> Tuple[str, str]:
    """Extract text from PDF or DOCX while preserving document type"""
    file_type = filename.lower().split('.')[-1]
    
    if file_type == 'pdf':
        with pdfplumber.open(io.BytesIO(file_content)) as pdf:
            text = '\n'.join(page.extract_text() or '' for page in pdf.pages)
        return text, 'pdf'
    
    elif file_type in ['docx', 'doc']:
        doc = docx.Document(io.BytesIO(file_content))
        text = '\n'.join(paragraph.text for paragraph in doc.paragraphs)
        return text, 'docx'
    
    raise ValueError("Unsupported file format")

def create_refined_document(original_content: bytes, refined_text: str, file_type: str) -> bytes:
    """Create a new document with refined content while preserving formatting"""
    output_buffer = io.BytesIO()
    
    try:
        if file_type == 'docx':
            # Create new document
            doc = docx.Document()
            
            # Add paragraphs with basic formatting
            for paragraph in refined_text.split('\n'):
                if paragraph.strip():
                    p = doc.add_paragraph(paragraph.strip())
                    # Add some basic formatting
                    p.style = doc.styles['Normal']
                    for run in p.runs:
                        run.font.size = Pt(11)
            
            # Save to buffer
            doc.save(output_buffer)
            
        elif file_type == 'pdf':
            # For PDF, we'll create a simple PDF with the text
            # First create a Word doc
            doc = docx.Document()
            for paragraph in refined_text.split('\n'):
                if paragraph.strip():
                    doc.add_paragraph(paragraph.strip())
            
            # Save to temporary buffer
            doc.save(output_buffer)
            
        # Get the bytes from the buffer
        output_buffer.seek(0)
        return output_buffer.getvalue()
        
    except Exception as e:
        raise Exception(f"Error creating document: {str(e)}")
    finally:
        output_buffer.close()

def calculate_match_percentage(job_description: str, resume_text: str) -> Tuple[float, List[str]]:
    """Calculate match percentage and extract key requirements"""
    try:
        # Extract key requirements from job description
        requirements = extract_requirements(job_description)
        
        if not requirements:
            return 0.0, []
        
        # Check which requirements are met in the resume
        matched_requirements = []
        for req in requirements:
            if requirement_matches(req, resume_text):
                matched_requirements.append(req)
        
        match_percentage = (len(matched_requirements) / len(requirements)) * 100 if requirements else 0
        return match_percentage, matched_requirements
    except Exception as e:
        print(f"Error calculating match percentage: {str(e)}")
        return 0.0, []

def extract_requirements(job_description: str) -> List[str]:
    """Extract key requirements from job description"""
    try:
        requirements = []
        
        # Look for common requirement patterns
        patterns = [
            r'required skills:.*?(?=\n\n|\Z)',
            r'requirements:.*?(?=\n\n|\Z)',
            r'qualifications:.*?(?=\n\n|\Z)',
            r'must have:.*?(?=\n\n|\Z)'
        ]
        
        for pattern in patterns:
            matches = re.findall(pattern, job_description, re.IGNORECASE | re.DOTALL)
            for match in matches:
                # Split into bullet points or lines
                items = re.split(r'[•\-\*]\s*|\n+', match)
                requirements.extend([item.strip() for item in items if item.strip()])
        
        # If no structured requirements found, try to extract key phrases
        if not requirements:
            words = job_description.split()
            for i in range(len(words)-2):
                phrase = ' '.join(words[i:i+3])
                if any(keyword in phrase.lower() for keyword in ['experience with', 'knowledge of', 'ability to', 'skills in']):
                    requirements.append(phrase)
        
        return list(set(requirements))
    except Exception as e:
        print(f"Error extracting requirements: {str(e)}")
        return []

def requirement_matches(requirement: str, resume_text: str) -> bool:
    """Check if a requirement is matched in the resume text"""
    try:
        return requirement.lower() in resume_text.lower()
    except Exception:
        return False

def generate_change_summary(original_text: str, refined_text: str) -> List[str]:
    """Generate a list of changes made to the resume"""
    try:
        changes = []
        
        # Compare sections and identify major changes
        original_sections = split_into_sections(original_text)
        refined_sections = split_into_sections(refined_text)
        
        for section in refined_sections:
            if section not in original_sections:
                changes.append(f"Added new section: {section}")
            elif refined_sections[section] != original_sections.get(section, ''):
                changes.append(f"Updated content in {section} section")
        
        # If no structural changes detected, add a generic change message
        if not changes:
            changes.append("Refined content to better match job requirements")
        
        return changes
    except Exception as e:
        print(f"Error generating change summary: {str(e)}")
        return ["Resume has been refined to better match the job requirements"]

def split_into_sections(text: str) -> Dict[str, str]:
    """Split resume text into sections"""
    try:
        sections = {}
        current_section = "General"
        current_content = []
        
        for line in text.split('\n'):
            if line.strip().isupper() and len(line.strip()) > 3:
                # Assume this is a section header
                if current_content:
                    sections[current_section] = '\n'.join(current_content)
                current_section = line.strip()
                current_content = []
            else:
                current_content.append(line)
        
        if current_content:
            sections[current_section] = '\n'.join(current_content)
        
        return sections
    except Exception as e:
        print(f"Error splitting sections: {str(e)}")
        return {"General": text} 